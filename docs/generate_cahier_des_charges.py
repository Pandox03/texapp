# -*- coding: utf-8 -*-
"""Generate AbrajTex cahier des charges (.docx)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent / "Cahier_des_charges_AbrajTex.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title_page(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CAHIER DES CHARGES")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Application de gestion\nimport textile & ventes en gros")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

    doc.add_paragraph()

    app = doc.add_paragraph()
    app.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = app.add_run("AbrajTex")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    doc.add_paragraph()

    client = doc.add_paragraph()
    client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = client.add_run("ABRAJE TEX")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Version 1.0\n{date.today().strftime('%d/%m/%Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Table des matières", level=1)
    sections = [
        "1. Introduction et contexte",
        "2. Présentation du projet",
        "3. Objectifs du système",
        "4. Périmètre et utilisateurs",
        "5. Exigences fonctionnelles",
        "6. Exigences non fonctionnelles",
        "7. Architecture technique",
        "8. Contraintes, hypothèses et dépendances",
        "9. Livrables et critères d'acceptation",
        "10. Glossaire",
        "11. Annexes",
    ]
    for item in sections:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.left_indent = Cm(0.5)
    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_shading(hdr[i], "E8F4F3")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_page(doc)
    add_toc(doc)

    # 1. Introduction
    doc.add_heading("1. Introduction et contexte", level=1)
    doc.add_paragraph(
        "Le présent cahier des charges définit les besoins fonctionnels et techniques "
        "de l'application AbrajTex, développée pour ABRAJE TEX, entreprise marocaine "
        "spécialisée dans l'importation de tissus (principalement depuis la Chine) "
        "et la vente en gros de rouleaux de tissu à des clients professionnels."
    )
    doc.add_paragraph(
        "L'activité repose sur la réception de conteneurs, la gestion du stock par type "
        "de tissu et par rouleau, l'enregistrement des ventes au m², la facturation "
        "conforme à la réglementation marocaine (HT / TVA / TTC), le suivi des paiements "
        "clients et la gestion des créances historiques (crédits legacy)."
    )
    doc.add_paragraph(
        "Ce document sert de référence pour la maintenance, l'évolution du produit "
        "et la validation des livrables entre le commanditaire et l'équipe technique."
    )

    # 2. Présentation
    doc.add_heading("2. Présentation du projet", level=1)
    doc.add_heading("2.1 Identité du projet", level=2)
    add_table(
        doc,
        ["Élément", "Description"],
        [
            ["Nom du projet", "AbrajTex"],
            ["Client / bénéficiaire", "ABRAJE TEX"],
            ["Type de solution", "Application web full-stack (SPA + API REST)"],
            ["URL de production", "https://abrajtex.com"],
            ["Devise", "Dirham marocain (MAD)"],
            ["Langues de l'interface", "Français et Arabe (bilingue)"],
            ["Dépôt source", "https://github.com/Pandox03/AbrajTex"],
        ],
    )

    doc.add_heading("2.2 Contexte métier", level=2)
    add_bullets(
        doc,
        [
            "Réception et suivi des conteneurs d'importation (origine Chine par défaut).",
            "Déclinaison du stock par type de tissu, couleur et rouleau individuel.",
            "Vente de rouleaux avec quantité en m² et prix unitaire négocié.",
            "Émission de factures PDF avec ventilation HT, TVA (20 %) et TTC.",
            "Encaissement partiel ou total des clients avec preuve de paiement.",
            "Gestion séparée des ventes stock et des crédits historiques (legacy).",
            "Tableaux de bord adaptés aux rôles administrateur, secrétaire et comptable.",
        ],
    )

    # 3. Objectifs
    doc.add_heading("3. Objectifs du système", level=1)
    doc.add_heading("3.1 Objectifs principaux", level=2)
    add_numbered(
        doc,
        [
            "Centraliser la traçabilité des conteneurs, du stock et des rouleaux vendus.",
            "Automatiser la facturation et le calcul fiscal (TVA incluse dans les montants saisis).",
            "Offrir une vision claire des soldes clients (ventes stock, crédits, total).",
            "Réduire les erreurs manuelles lors des ventes et des paiements.",
            "Permettre un contrôle d'accès par rôle (administration, opérations, finance).",
            "Fournir un historique d'activité pour l'audit et la conformité interne.",
        ],
    )

    doc.add_heading("3.2 Objectifs secondaires (IA)", level=2)
    add_bullets(
        doc,
        [
            "Assistant conversationnel bilingue (FR/AR) pour interroger les données métier.",
            "Résumé automatique du profil client (historique, soldes, risques).",
            "Suggestions de prix intelligentes lors de la création d'une vente.",
            "Enrichissement des conteneurs et types de tissu avec coûts et notes marché.",
        ],
    )

    # 4. Périmètre
    doc.add_heading("4. Périmètre et utilisateurs", level=1)
    doc.add_heading("4.1 Utilisateurs cibles", level=2)
    add_table(
        doc,
        ["Rôle", "Profil", "Accès principal"],
        [
            ["Administrateur", "Dirigeant / responsable IT", "Accès complet : utilisateurs, logs, suppression, dashboard global"],
            ["Secrétaire", "Opérateur commercial / logistique", "Conteneurs, stock, ventes, clients, factures, paiements (création)"],
            ["Comptable", "Responsable finance", "Lecture factures, paiements, profils clients, dashboard comptable"],
        ],
    )

    doc.add_heading("4.2 Matrice des droits (synthèse)", level=2)
    add_table(
        doc,
        ["Module", "Admin", "Secrétaire", "Comptable"],
        [
            ["Dashboard global", "Oui", "Non (dashboard secrétaire)", "Non (dashboard comptable)"],
            ["Conteneurs & stock", "Oui", "Oui", "Non"],
            ["Types de tissu", "Oui (CRUD complet)", "Oui (CRU)", "Non"],
            ["Ventes & crédits", "Oui", "Oui", "Non"],
            ["Clients", "Oui (CRUD)", "Oui (CRU)", "Lecture"],
            ["Factures (PDF)", "Oui", "Oui (création/édition)", "Lecture"],
            ["Paiements", "Oui", "Oui (création)", "Lecture"],
            ["Utilisateurs", "Oui", "Non", "Non"],
            ["Journal d'activité", "Oui", "Non", "Non"],
            ["Assistant IA", "Oui", "Oui (sans pricing hint)", "Oui (chat + résumé client)"],
        ],
    )

    doc.add_heading("4.3 Hors périmètre", level=2)
    add_bullets(
        doc,
        [
            "Comptabilité générale complète (balance, grand livre, clôture annuelle).",
            "Gestion de la paie et des ressources humaines.",
            "E-commerce B2C et boutique en ligne.",
            "Intégration ERP externe (SAP, Sage, etc.) — non prévue en v1.",
            "Application mobile native (iOS / Android).",
        ],
    )

    # 5. Exigences fonctionnelles
    doc.add_heading("5. Exigences fonctionnelles", level=1)

    modules = [
        (
            "5.1 Authentification et profil",
            [
                "Connexion sécurisée par e-mail et mot de passe (token Sanctum).",
                "Déconnexion et invalidation du token.",
                "Consultation et modification du profil utilisateur connecté.",
                "Redirection automatique vers le tableau de bord selon le rôle.",
            ],
        ),
        (
            "5.2 Gestion des conteneurs",
            [
                "CRUD conteneurs : référence, date d'arrivée, origine, fournisseur, statut, notes.",
                "Statuts : en transit, arrivé, en traitement, clôturé.",
                "Ajout d'articles par conteneur (type de tissu, code couleur, quantité m², rouleaux estimés).",
                "Suivi des coûts d'achat, transport, douane, frais divers et notes marché (évolution v1.1).",
                "Vue détaillée conteneur avec liste des articles et rouleaux associés.",
            ],
        ),
        (
            "5.3 Types de tissu",
            [
                "Catalogue hiérarchique (parent / sous-types).",
                "Attributs : code, composition, largeur par défaut (cm), GSM, description.",
                "Champs coût unitaire et notes marché pour aide à la tarification.",
                "Suppression réservée à l'administrateur.",
            ],
        ),
        (
            "5.4 Stock et rouleaux",
            [
                "Stock global agrégé par type de tissu et couleur.",
                "Liste des rouleaux : numéro, dimensions, poids, statut (disponible, réservé, vendu).",
                "Création manuelle de rouleaux liés à un conteneur.",
                "Vérification de disponibilité lors d'une vente (stock temps réel).",
            ],
        ),
        (
            "5.5 Clients",
            [
                "Fiche client : nom, téléphone, e-mail, adresse, ICE, notes, délai de paiement.",
                "Liste complète des clients (pagination étendue pour usage interne).",
                "Profil client avec onglets : Ventes | Crédits | Paiements | Factures.",
                "Soldes distincts : solde ventes stock, solde crédits legacy, solde total.",
                "Suppression client réservée à l'administrateur.",
            ],
        ),
        (
            "5.6 Ventes (stock)",
            [
                "Création de vente : client, date, sélection de rouleaux disponibles.",
                "Saisie du prix unitaire et quantité m² par ligne ; calcul automatique du total.",
                "Génération d'une référence vente unique.",
                "Modification et suppression de vente (avec contrôles de cohérence stock).",
                "Suggestion de prix IA basée sur l'historique client et le contexte marché.",
            ],
        ),
        (
            "5.7 Crédits legacy",
            [
                "Enregistrement de créances historiques indépendantes du stock.",
                "Montant TTC, date, notes ; pas de lien obligatoire aux rouleaux.",
                "Édition et suppression (avec suppression des paiements et factures liés).",
                "Paiements ciblés sur un crédit via sale_id ; affichage expandable sous chaque crédit.",
                "Exclusion des paiements crédit de l'onglet Paiements principal.",
            ],
        ),
        (
            "5.8 Facturation",
            [
                "Création de facture à partir d'une vente (montant partiel ou total restant).",
                "Numérotation automatique des factures.",
                "Ventilation HT / TVA 20 % / TTC à partir du montant TTC saisi.",
                "Génération et téléchargement PDF (logo entreprise intégré).",
                "Notification e-mail aux administrateurs à l'enregistrement.",
                "Statut simplifié : envoyée (sent).",
            ],
        ),
        (
            "5.9 Paiements",
            [
                "Enregistrement paiement : client, montant, date, mode, référence, notes.",
                "Upload preuve de paiement (fichier joint).",
                "Allocation FIFO automatique sur les ventes stock uniquement.",
                "Paiements crédit : allocation directe via sale_id.",
                "Consultation preuve PDF/image par les rôles autorisés.",
                "Notification e-mail aux administrateurs à l'enregistrement.",
            ],
        ),
        (
            "5.10 Tableaux de bord",
            [
                "Dashboard admin : indicateurs globaux (ventes, encaissements, stock, clients).",
                "Dashboard secrétaire : focus opérationnel (conteneurs récents, ventes, alertes stock).",
                "Dashboard comptable : factures, paiements, soldes clients.",
            ],
        ),
        (
            "5.11 Utilisateurs et audit",
            [
                "CRUD utilisateurs (admin) : nom, e-mail, rôle, mot de passe.",
                "Protection du dernier compte administrateur (non supprimable / non rétrogradable).",
                "Journal d'activité : connexions, créations, modifications significatives.",
            ],
        ),
        (
            "5.12 Assistant IA (OpenRouter)",
            [
                "Chat flottant bilingue FR/AR accessible depuis toute l'application.",
                "Contexte métier injecté (stock, clients, ventes) pour réponses pertinentes.",
                "Résumé IA du profil client (GET /api/ai/clients/{id}/summary).",
                "Indication de prix (POST /api/ai/pricing-hint) sur page nouvelle vente.",
                "Modèle : Google Gemini 2.5 Flash via OpenRouter ; retries et timeouts configurés.",
            ],
        ),
    ]

    for title, items in modules:
        doc.add_heading(title, level=2)
        add_bullets(doc, items)

    doc.add_heading("5.13 Règles métier clés", level=2)
    add_table(
        doc,
        ["Règle", "Description"],
        [
            ["TVA", "Taux par défaut 20 % ; montants saisis en TTC, décomposition HT/TVA automatique"],
            ["FIFO paiements", "Paiements non ciblés affectés aux ventes stock par ordre chronologique"],
            ["Paiements crédit", "Liés explicitement à une vente crédit ; exclus du FIFO stock"],
            ["Soldes clients", "Solde ventes + solde crédits = solde total affiché"],
            ["Rouleau vendu", "Statut 'sold' + lien sale_id ; non sélectionnable pour nouvelle vente"],
            ["Facture partielle", "Montant facturé ≤ reste à facturer sur la vente"],
        ],
    )

    # 6. Non fonctionnel
    doc.add_heading("6. Exigences non fonctionnelles", level=1)
    doc.add_heading("6.1 Performance et disponibilité", level=2)
    add_bullets(
        doc,
        [
            "Temps de réponse API < 2 s pour les opérations courantes (hors génération PDF).",
            "Application accessible 24h/24 via hébergement web production.",
            "Build frontend optimisé (Vite) pour chargement initial < 5 s sur connexion standard.",
        ],
    )

    doc.add_heading("6.2 Sécurité", level=2)
    add_bullets(
        doc,
        [
            "Authentification API par token Bearer (Laravel Sanctum).",
            "Contrôle d'accès par middleware de rôle sur chaque route.",
            "Variables sensibles (.env) hors dépôt Git ; clé OpenRouter en production.",
            "APP_DEBUG=false en production ; HTTPS obligatoire.",
            "Validation des entrées côté serveur (Form Requests Laravel).",
        ],
    )

    doc.add_heading("6.3 Ergonomie et accessibilité", level=2)
    add_bullets(
        doc,
        [
            "Interface responsive (desktop prioritaire, tablette supportée).",
            "Bascule FR / AR avec traductions centralisées (locales).",
            "Navigation latérale persistante ; fil d'Ariane implicite par module.",
            "Messages d'erreur et confirmations en langue de l'interface.",
            "Support RTL pour l'arabe.",
        ],
    )

    doc.add_heading("6.4 Sauvegarde et maintenance", level=2)
    add_bullets(
        doc,
        [
            "Base MySQL 8 avec sauvegardes régulières côté hébergeur.",
            "Migrations Laravel versionnées pour évolutions schéma.",
            "Logs applicatifs Laravel pour diagnostic (storage/logs).",
            "Déploiement documenté (git pull, cache, build frontend, upload dist).",
        ],
    )

    # 7. Architecture
    doc.add_heading("7. Architecture technique", level=1)
    doc.add_heading("7.1 Stack technologique", level=2)
    add_table(
        doc,
        ["Couche", "Technologie", "Version / remarque"],
        [
            ["Backend API", "Laravel + Sanctum", "Laravel 12"],
            ["Frontend", "React + Vite + TypeScript", "React 19"],
            ["Styles", "Tailwind CSS", "v4"],
            ["Base de données", "MySQL", "8.x"],
            ["PDF factures", "DomPDF / Blade templates", "Logo : public/images/logo.png"],
            ["IA", "OpenRouter API", "google/gemini-2.5-flash"],
            ["Hébergement", "Hostinger", "SSH, PHP, MySQL, public_html"],
        ],
    )

    doc.add_heading("7.2 Modèle de données (entités principales)", level=2)
    add_table(
        doc,
        ["Entité", "Relations clés"],
        [
            ["containers", "1-N container_items, 1-N fabric_rolls"],
            ["fabric_types", "Hiérarchie parent_id ; 1-N rouleaux, articles"],
            ["fabric_rolls", "BelongsTo container, fabric_type, sale (optionnel)"],
            ["clients", "1-N sales, invoices, payments"],
            ["sales", "1-N sale_items ; flag legacy credit ; 1-N invoices, payments"],
            ["sale_items", "Lien vente ↔ rouleau + prix + m²"],
            ["invoices", "BelongsTo sale, client ; PDF généré à la demande"],
            ["payments", "BelongsTo client ; sale_id optionnel ; preuve fichier"],
            ["users", "Rôle enum : admin, secretaire, comptable"],
            ["activity_logs", "Audit trail des actions utilisateurs"],
        ],
    )

    doc.add_heading("7.3 Architecture déploiement", level=2)
    doc.add_paragraph(
        "En production, le frontend compilé (frontend/dist) est servi depuis public_html "
        "tandis que l'API Laravel est déployée sur le même serveur ou sous-domaine API. "
        "La variable VITE_API_URL pointe vers l'endpoint /api. Les fichiers uploadés "
        "(preuves de paiement) sont stockés via storage Laravel avec lien symbolique public."
    )

    doc.add_heading("7.4 Services métier principaux", level=2)
    add_bullets(
        doc,
        [
            "BillingService — calcul TVA, FIFO, soldes, facturation.",
            "AdminNotificationService — e-mails factures et paiements aux admins.",
            "AiService / OpenRouterService — chat, résumés, pricing hints.",
            "PricingContextService — agrégation contexte prix (historique, coûts, marché).",
        ],
    )

    # 8. Contraintes
    doc.add_heading("8. Contraintes, hypothèses et dépendances", level=1)
    doc.add_heading("8.1 Contraintes", level=2)
    add_bullets(
        doc,
        [
            "Devise unique : MAD.",
            "TVA marocaine standard à 20 % (configurable dans BillingService).",
            "Un seul entrepôt logique (stock global, pas de multi-sites).",
            "Connexion Internet requise (application web, API IA externe).",
        ],
    )

    doc.add_heading("8.2 Hypothèses", level=2)
    add_bullets(
        doc,
        [
            "Les utilisateurs disposent d'un navigateur moderne (Chrome, Edge, Firefox).",
            "Les rouleaux sont identifiés de façon unique par type + numéro de rouleau.",
            "Les montants commerciaux sont exprimés TTC.",
            "Le compte administrateur initial est créé via seeder (.env ADMIN_EMAIL/PASSWORD).",
        ],
    )

    doc.add_heading("8.3 Dépendances externes", level=2)
    add_bullets(
        doc,
        [
            "OpenRouter API (clé OPENROUTER_API_KEY) pour fonctionnalités IA.",
            "Serveur SMTP ou service mail pour notifications administrateur.",
            "Hébergeur Hostinger (PHP 8.x, MySQL, certificat SSL).",
        ],
    )

    # 9. Livrables
    doc.add_heading("9. Livrables et critères d'acceptation", level=1)
    doc.add_heading("9.1 Livrables", level=2)
    add_numbered(
        doc,
        [
            "Code source backend (Laravel) et frontend (React) versionné sur GitHub.",
            "Base de données migrée et seedée (environnement de production).",
            "Application déployée et accessible sur https://abrajtex.com.",
            "Documentation README (installation, déploiement, structure projet).",
            "Guide déploiement Hostinger (deploy/HOSTINGER.md).",
            "Présent cahier des charges.",
        ],
    )

    doc.add_heading("9.2 Critères d'acceptation (extrait)", level=2)
    add_table(
        doc,
        ["ID", "Critère", "Validation"],
        [
            ["AC-01", "Un secrétaire peut créer une vente et décrémenter le stock", "Test manuel OK"],
            ["AC-02", "Une facture PDF téléchargeable reflète HT/TVA/TTC corrects", "Test manuel OK"],
            ["AC-03", "Un paiement non ciblé réduit le solde FIFO des ventes stock", "Test manuel OK"],
            ["AC-04", "Un paiement crédit n'affecte que le crédit concerné", "Test manuel OK"],
            ["AC-05", "Le comptable voit factures/paiements sans modifier le stock", "Test rôle OK"],
            ["AC-06", "L'interface bascule FR/AR sans rechargement complet", "Test UI OK"],
            ["AC-07", "L'assistant IA répond en FR ou AR selon la langue active", "Test IA (si clé configurée)"],
        ],
    )

    # 10. Glossaire
    doc.add_heading("10. Glossaire", level=1)
    add_table(
        doc,
        ["Terme", "Définition"],
        [
            ["m²", "Mètre carré — unité de vente des tissus"],
            ["Rouleau", "Unité physique de stock (fabric_roll) avec dimensions et statut"],
            ["Conteneur", "Arrivage d'import regroupant plusieurs articles/rouleaux"],
            ["Crédit legacy", "Créance historique non liée au stock actuel"],
            ["FIFO", "First In, First Out — affectation des paiements par ancienneté"],
            ["HT", "Hors taxes"],
            ["TTC", "Toutes taxes comprises"],
            ["TVA", "Taxe sur la valeur ajoutée (20 % par défaut)"],
            ["ICE", "Identifiant Commun de l'Entreprise (Maroc)"],
            ["GSM", "Grammage du tissu (g/m²)"],
            ["Sanctum", "Package Laravel d'authentification API par token"],
        ],
    )

    # 11. Annexes
    doc.add_heading("11. Annexes", level=1)
    doc.add_heading("11.1 Arborescence des écrans", level=2)
    add_table(
        doc,
        ["Route", "Écran", "Rôles"],
        [
            ["/", "Tableau de bord (selon rôle)", "Tous"],
            ["/containers", "Liste conteneurs", "Admin, Secrétaire"],
            ["/containers/:id", "Détail conteneur", "Admin, Secrétaire"],
            ["/stock", "Stock global", "Admin, Secrétaire"],
            ["/fabric-types", "Types de tissu", "Admin, Secrétaire"],
            ["/sales", "Liste ventes", "Admin, Secrétaire"],
            ["/sales/new", "Nouvelle vente", "Admin, Secrétaire"],
            ["/credits/new", "Nouveau crédit", "Admin, Secrétaire"],
            ["/clients", "Liste clients", "Admin, Secrétaire"],
            ["/clients/:id", "Profil client", "Admin, Secrétaire, Comptable"],
            ["/invoices", "Factures", "Admin, Secrétaire, Comptable"],
            ["/invoices/generer", "Générer facture", "Admin, Secrétaire"],
            ["/payments", "Paiements", "Admin, Comptable"],
            ["/users", "Utilisateurs", "Admin"],
            ["/logs", "Journal d'activité", "Admin"],
            ["/profile", "Mon profil", "Tous"],
        ],
    )

    doc.add_heading("11.2 Variables d'environnement essentielles", level=2)
    add_table(
        doc,
        ["Variable", "Description"],
        [
            ["APP_URL", "URL de l'API en production"],
            ["FRONTEND_URL", "URL du frontend (CORS, mails)"],
            ["DB_*", "Connexion MySQL"],
            ["ADMIN_EMAIL / ADMIN_PASSWORD", "Compte admin initial (seeder)"],
            ["VITE_API_URL", "URL API côté frontend (build)"],
            ["OPENROUTER_API_KEY", "Clé API pour assistant IA"],
            ["MAIL_*", "Configuration envoi e-mails"],
        ],
    )

    doc.add_heading("11.3 Historique des versions", level=2)
    add_table(
        doc,
        ["Version", "Date", "Évolutions majeures"],
        [
            ["1.0", "Juillet 2026", "MVP : conteneurs, stock, ventes, factures, paiements, rôles"],
            ["1.1", "Juillet–Août 2026", "Crédits legacy, soldes séparés, paiements crédit, UX profil client"],
            ["1.2", "Août 2026", "Assistant IA, coûts conteneurs, suggestions de prix (en déploiement)"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— Fin du document —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(f"Generated: {OUTPUT}")
